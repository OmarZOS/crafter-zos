



import re
from docx import Document

    
def load_document(filename):
    # load Word document
    doc = Document(filename)
    return doc


def save_document(doc,filename):
    # save the modified document
    doc.save(filename)
    
    
def replace_in_file(doc,old,new):
    # replace text
    
    for para in doc.paragraphs:
        # Loop through runs (style spans)
        for run in para.runs:
            # if there is text on this run, replace it
            if run.text:
                # get the replacement text
                # print(run.text)
                replaced_text = re.sub(old, new, run.text, 999)
                if replaced_text != run.text:
                    # if the replaced text is not the same as the original
                    # replace the text and increment the number of occurences
                    run.text = replaced_text
                    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        paragraph.text = paragraph.text.replace(old, new)
    
    
    return doc
    
    


if __name__=="__main__":
    doc = load_document("Template.docx")
    replace_in_file(doc, "old" ,"new")
    save_document(doc,"Template2.docx")
    
    
    